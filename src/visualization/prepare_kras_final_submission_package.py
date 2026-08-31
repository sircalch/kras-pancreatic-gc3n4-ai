"""
prepare_kras_final_submission_package.py
Packages all generated files for Article 3 (KRAS-G12D & g-C3N4) into an official
submission-ready folder and ZIP package.
"""

import os
import shutil
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def create_kras_cover_letter(sub_dir):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)
    
    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_after = Pt(14)
    p_h.add_run(
        "Andrés Monreal Hernández, Ph.D.\n"
        "Universidad Estatal de Sonora\n"
        "Hermosillo, Sonora, Mexico\n"
        "Email: andres.monreal@ues.mx | ORCID: 0009-0009-1207-8597\n"
        "Date: August 30, 2026\n"
    ).font.bold = True
    
    p_ed = doc.add_paragraph()
    p_ed.paragraph_format.space_after = Pt(12)
    p_ed.add_run(
        "To: The Editor-in-Chief\n"
        "Beilstein Journal of Nanotechnology\n"
        "Beilstein-Institut, Frankfurt am Main, Germany\n"
    )
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    p_sub.add_run("Subject: Submission of Original Research Article for Peer Review").font.bold = True
    
    doc.add_paragraph("Dear Editor-in-Chief,")
    doc.add_paragraph(
        "On behalf of my co-authors (Sara Lizbeth Franco Amaya, Carlos Ivanhoe Martínez Osorio, and myself), "
        "I am pleased to submit our original research manuscript titled:"
    )
    
    p_t = doc.add_paragraph()
    p_t.paragraph_format.left_indent = Inches(0.4)
    p_t.paragraph_format.space_after = Pt(10)
    r_t = p_t.add_run("“Quantum-Informed and Machine Learning QSAR Investigation of Graphitic Carbon Nitride (g-C3N4) Nanocarriers Delivering Allosteric Inhibitors Targeting Oncogenic KRAS-G12D in Pancreatic Ductal Adenocarcinoma”")
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0, 105, 92)
    
    doc.add_paragraph(
        "for consideration for publication as a Full Research Article in the Beilstein Journal of Nanotechnology."
    )
    
    doc.add_paragraph(
        "This study represents the first comprehensive computational and Explainable AI framework evaluating 2D polymeric graphitic carbon nitride (g-C3N4) "
        "and B/P-doped nanostructures for the targeted delivery of 36 KRAS-G12D allosteric inhibitors and pancreatic therapeutics against crystal structure PDB 7RPZ."
    )
    
    doc.add_paragraph(
        "All authors have approved the manuscript and confirm no competing interests."
    )
    
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(14)
    p_sign.add_run(
        "Sincerely,\n\n"
        "Andrés Monreal Hernández, Ph.D. (Corresponding Author)\n"
        "Universidad Estatal de Sonora, Mexico\n"
        "Email: andres.monreal@ues.mx"
    )
    
    out_docx = os.path.join(sub_dir, "01_Cover_Letter_Beilstein_KRAS.docx")
    doc.save(out_docx)
    print(f"Generated KRAS Cover Letter: {out_docx}")

def build_kras_submission_bundle():
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    sub_dir = os.path.join(base_dir, "manuscript", "submission_ready")
    os.makedirs(sub_dir, exist_ok=True)
    
    create_kras_cover_letter(sub_dir)
    
    manuscripts = {
        os.path.join(base_dir, "manuscript", "KRAS_gC3N4_Full_Q1_Research_Paper_Monreal_Hernandez_et_al.docx"): "02_Manuscript_KRAS_gC3N4_Full_Q1_Research_Paper.docx",
        os.path.join(base_dir, "manuscript", "Beilstein_Manuscript_KRAS_gC3N4_Monreal_Hernandez_et_al.docx"): "02_Manuscript_KRAS_gC3N4_Beilstein_Submission.docx"
    }
    
    for src_docx, dst_name in manuscripts.items():
        dst_docx = os.path.join(sub_dir, dst_name)
        if os.path.exists(src_docx):
            shutil.copyfile(src_docx, dst_docx)
        
    zip_path = os.path.join(base_dir, "kras-pancreatic-gC3N4-ai-FINAL-SUBMISSION-READY.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_f:
        for root, dirs, files in os.walk(sub_dir):
            for file in files:
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, sub_dir)
                zip_f.write(file_path, os.path.join("submission_ready", rel_path))
                
    print(f"\n=======================================================")
    print(f">>> KRAS SUBMISSION PACKAGE GENERATED SUCCESSFULLY ({os.path.getsize(zip_path)} bytes) <<<")
    print(f" -> {zip_path}")
    print(f"=======================================================")

if __name__ == "__main__":
    build_kras_submission_bundle()
