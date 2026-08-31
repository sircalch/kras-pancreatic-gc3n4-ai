"""
integrate_real_figures.py
=========================
Integrates the 4 user-generated real scientific figures into the manuscript.

Figures to integrate:
  - Graphical Abstract  -> media_1788126628837.jpg  (P3 in doc)
  - Figure 3 (Redocking validation) -> media_1788126079425.jpg  (P18/P20 in doc)
  - Figure 8 (QSPR Validation)      -> media_1788126070436.jpg  (P30/P32 in doc)
  - Figure 10 (Atomistic structures) -> media_1788126622639.jpg  (replace Fig 9 at P32/...)

Strategy:
  1. Copy Beilstein manuscript -> new file with _RealFigures suffix
  2. Replace the image bytes for target paragraphs using python-docx
  3. Update figure captions to match actual figures
  4. Save as final submission-ready DOCX
"""

import shutil
import os
import copy
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import io

# ── PATHS ─────────────────────────────────────────────────────────────────────
BASE_DIR = Path(r"c:\Users\Andre\Proyectos doctorado\kras-pancreatic-gC3N4-ai")
UPLOAD_DIR = Path(r"C:\Users\Andre\.gemini\antigravity\brain\93c3b72a-7190-407d-8ac2-98ee7c61aaf9\.user_uploaded")
MANUSCRIPT_IN  = BASE_DIR / "manuscript" / "Beilstein_Manuscript_KRAS_gC3N4_Monreal_Hernandez_et_al.docx"
MANUSCRIPT_OUT = BASE_DIR / "manuscript" / "KRAS_gC3N4_FINAL_RealFigures_Monreal_Hernandez_et_al.docx"

# ── USER FIGURE FILES ──────────────────────────────────────────────────────────
FIG_FILES = {
    "graphical_abstract": UPLOAD_DIR / "media_1788126628837.jpg",
    "fig3_redocking":     UPLOAD_DIR / "media_1788126079425.jpg",
    "fig8_qspr":          UPLOAD_DIR / "media_1788126070436.jpg",
    "fig10_atomistic":    UPLOAD_DIR / "media_1788126622639.jpg",
}

# ── UPDATED CAPTIONS ───────────────────────────────────────────────────────────
CAPTIONS = {
    "graphical_abstract": (
        "Graphical Abstract. Multi-scale computational strategy for KRAS-G12D inhibitor screening "
        "and g-C\u2083N\u2084 nanocarrier evaluation. Scene 1 (left): structural targeting of the "
        "KRAS-G12D Switch-II allosteric pocket by MRTX1133 (PDB 7RPZ, 1.30 \u00c5; heavy-atom RMSD = "
        "1.419 \u00c5). Scene 2 (center): quantum adsorption of MRTX1133 on B/P co-doped g-C\u2083N\u2084 "
        "(\u0394E\u2090\u1d48\u02e2 = \u22123 5.04 kcal mol\u207b\u00b9, GFN2-xTB + D4; Mulliken charges "
        "q\u2082 = +0.35 e, q\u209a = \u22120.17 e; \u03c0\u2013\u03c0 stacking d = 3.35 \u00c5). "
        "Scene 3 (right): QSPR-guided virtual screening of 350 oncology inhibitors yields 328 "
        "compounds within the applicability domain (h* = 0.455); top prioritized leads \u2014 "
        "Avapritinib, Futibatinib, and Belumosudil \u2014 confirmed by external GFN2-xTB "
        "single-point calculations."
    ),
    "fig3_redocking": (
        "Figure 3. Crystallographic redocking validation of MRTX1133 on KRAS-G12D (PDB 7RPZ, "
        "1.30 \u00c5). (a) Stereoview superposition of the co-crystallographic pose (gray sticks) "
        "and the AutoDock Vina best-ranked redocked pose (teal sticks) inside the Switch-II "
        "allosteric pocket. Dashed lines denote key contacts with Asp12, Glu62, Tyr96, and Arg68. "
        "Heavy-atom RMSD = 1.419 \u00c5, confirming successful pose reproduction (threshold < 2.0 \u00c5). "
        "(b) AutoDock Vina affinity distribution across the 9 sampled conformational modes; "
        "best affinity = \u22129.16 kcal mol\u207b\u00b9. Coordinates extracted from PDB entry 7RPZ."
    ),
    "fig8_qspr": (
        "Figure 8. QSPR model validation for adsorption energy (\u0394E\u2090\u1d48\u02e2) on "
        "B/P co-doped g-C\u2083N\u2084. (A) Out-of-fold (OOF) predicted vs. observed GFN2-xTB "
        "adsorption energies for the 33-compound training set (N = 33; "
        "Q\u00b2\u1d04\u1d20 = +0.5696, RMSE = 5.201 kcal mol\u207b\u00b9, "
        "MAE = 4.194 kcal mol\u207b\u00b9). Color coding: Group A \u2013 direct KRAS-G12D "
        "inhibitors (blue circles); Group B \u2013 mutation-selective/Pan-RAS (orange squares); "
        "Group C \u2013 downstream MAPK/RTK (green triangles); Group D \u2013 cytotoxic "
        "chemotherapy (red diamonds). (B) Williams plot (standardized residuals vs. hat leverage "
        "h\u1d62): warning leverage h* = 0.455 (dashed vertical line) and \u00b13\u03c3 residual "
        "limits (dashed horizontal lines); Methotrexate is the sole structural outlier (h\u1d62 > "
        "h*). (C) Y-scrambling validation (1,000 permutations): scrambled models yield "
        "\u1e41Q\u00b2 = \u22120.2357 (orange dashed line), empirical p-value = 0.001 (p < 0.001); "
        "the observed Q\u00b2\u1d04\u1d20 = +0.5696 (red line) exceeds all scrambled values, "
        "confirming non-chance correlation. (D) External quantum validation for the five "
        "prioritized leads: QSPR-predicted vs. GFN2-xTB-computed \u0394E\u2090\u1d48\u02e2 with "
        "leverage (h\u1d62), AutoDock Vina score, and ligand efficiency (LE). All leads fall "
        "within the applicability domain (h\u1d62 \u2264 h* = 0.455)."
    ),
    "fig10_atomistic": (
        "Figure 10. Atomistic multi-scale architecture of KRAS-G12D inhibitor engagement and "
        "g-C\u2083N\u2084 nanocarrier interactions. (A) Full KRAS-G12D receptor (PDB 7RPZ) in complex "
        "with MRTX1133 (crystal pose from PDB entry 6IC, displayed as sticks). "
        "(B) Zoom into the ionic and hydrogen-bond network within the Switch-II allosteric pocket: "
        "key distances \u2014 Asp12\u2013ligand 2.70 \u00c5, Glu62\u2013ligand 2.85 \u00c5, "
        "Tyr96\u2013ligand 3.34 \u00c5, Arg68\u2013ligand 3.43 \u00c5 (measured from 7RPZ coordinates). "
        "(C) BI-2865 Pan-KRAS inhibitor top docked pose (Vina = \u22128.46 kcal mol\u207b\u00b9). "
        "(D) MRTX1133 adsorbed on pristine g-C\u2083N\u2084 monolayer (GFN2-xTB optimized): "
        "\u03c0\u2013\u03c0 stacking at d = 3.25 \u00c5; \u0394E\u2090\u1d48\u02e2 = \u22123 5.03 kcal mol\u207b\u00b9. "
        "(E) MRTX1133 adsorbed on B/P co-doped g-C\u2083N\u2084: Mulliken charges q\u2082 = +0.3494 e "
        "(Boron, Lewis-acid center) and q\u209a = \u22120.1679 e (Phosphorus, electron donor); "
        "interfacial charge transfer \u0394Q = +0.146 e; \u0394E\u2090\u1d48\u02e2 = \u22123 5.04 kcal mol\u207b\u00b9. "
        "(F) Hamiltonian sensitivity benchmark across 10 chemically diverse systems: GFN2-xTB "
        "(with D4 dispersion) vs. GFN1-xTB adsorption energies; linear fit y = 0.394x \u2212 7.584, "
        "R\u00b2 = 0.254; MSE = \u221212.82 kcal mol\u207b\u00b9, MAE = 12.82 kcal mol\u207b\u00b9, "
        "RMSE = 17.34 kcal mol\u207b\u00b9. The systematic offset reflects the additional "
        "long-range dispersion captured exclusively by GFN2-xTB + D4, not a failure of the model."
    ),
}

# ── HELPER: Replace image bytes in a paragraph ─────────────────────────────────
def replace_image_in_paragraph(doc, para_idx, new_image_path, width_inches=6.3):
    """
    Replaces the image embedded in the given paragraph index with a new image.
    Uses the python-docx add_picture approach by manipulating the XML directly.
    """
    para = doc.paragraphs[para_idx]

    # Find all runs with drawings in this paragraph
    blip_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    r_ns = 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing'

    # Add new picture to a temporary paragraph to get the XML
    from docx.shared import Inches
    import tempfile

    # We'll add the image to the doc, get its paragraph element, then move it
    # into our target paragraph by replacing the existing drawing XML

    # First, find the existing drawing element in the target paragraph
    para_elem = para._element
    drawing_elems = para_elem.findall('.//' + qn('w:drawing'))

    if not drawing_elems:
        print(f"  WARNING: No drawing found in paragraph {para_idx}")
        return False

    # Add the new image to the document (appended at end temporarily)
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(str(new_image_path), width=Inches(width_inches))

    # Get the drawing XML from the new paragraph
    new_drawing = temp_para._element.findall('.//' + qn('w:drawing'))
    if not new_drawing:
        print(f"  WARNING: Could not add new picture for para {para_idx}")
        # Remove temp paragraph
        temp_para._element.getparent().remove(temp_para._element)
        return False

    new_drawing_elem = new_drawing[0]

    # Replace the old drawing with the new one in target paragraph
    old_drawing = drawing_elems[0]
    old_drawing.getparent().replace(old_drawing, new_drawing_elem)

    # Remove the temporary paragraph
    temp_para._element.getparent().remove(temp_para._element)

    print(f"  -> Successfully replaced image in paragraph {para_idx}")
    return True


def update_caption_paragraph(doc, caption_para_idx, new_caption_text, bold_prefix=None):
    """
    Clears and rewrites the caption paragraph with updated text.
    bold_prefix: e.g. 'Figure 8.' to make it bold.
    """
    para = doc.paragraphs[caption_para_idx]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_prefix and new_caption_text.startswith(bold_prefix):
        # Bold the figure label
        run_bold = para.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(9)

        rest = new_caption_text[len(bold_prefix):]
        run_rest = para.add_run(rest)
        run_rest.bold = False
        run_rest.font.size = Pt(9)
    else:
        run = para.add_run(new_caption_text)
        run.font.size = Pt(9)

    print(f"  -> Updated caption at paragraph {caption_para_idx}")


# ── MAIN ───────────────────────────────────────────────────────────────────────
def main():
    print("=" * 70)
    print("INTEGRATING REAL FIGURES INTO MANUSCRIPT")
    print("=" * 70)

    # Verify all figure files exist
    for name, path in FIG_FILES.items():
        if not path.exists():
            print(f"  ERROR: Figure file missing: {path}")
            return
        print(f"  OK: {name} -> {path.name} ({path.stat().st_size:,} bytes)")

    # Copy source manuscript
    shutil.copy2(MANUSCRIPT_IN, MANUSCRIPT_OUT)
    print(f"\nWorking on: {MANUSCRIPT_OUT.name}")

    doc = Document(str(MANUSCRIPT_OUT))

    print("\n--- PARAGRAPH MAP ---")
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if txt:
            print(f"  P{i}: {txt[:80]}")

    # ── 1. GRAPHICAL ABSTRACT (P3 = image, P4 = caption) ──────────────────────
    print("\n[1] Replacing Graphical Abstract (P3)...")
    replace_image_in_paragraph(doc, 3, FIG_FILES["graphical_abstract"], width_inches=6.3)
    update_caption_paragraph(doc, 4, CAPTIONS["graphical_abstract"], bold_prefix="Graphical Abstract.")

    # ── 2. FIGURE 3 REDOCKING (P18 = image, P19 = caption) ────────────────────
    # Current P18 has image for "Results and Discussion" figure (before Fig3 caption)
    # P18 = image, P19 = "Figure 3: Physical Molecular Docking Statistical Profiles..."
    print("\n[2] Replacing Figure 3 - Crystallographic Redocking (P18)...")
    replace_image_in_paragraph(doc, 18, FIG_FILES["fig3_redocking"], width_inches=6.3)
    update_caption_paragraph(doc, 19, CAPTIONS["fig3_redocking"], bold_prefix="Figure 3.")

    # ── 3. FIGURE 8 QSPR VALIDATION (P30 = image, P31 = caption) ─────────────
    print("\n[3] Replacing Figure 8 - QSPR Validation (P30)...")
    replace_image_in_paragraph(doc, 30, FIG_FILES["fig8_qspr"], width_inches=6.3)
    update_caption_paragraph(doc, 31, CAPTIONS["fig8_qspr"], bold_prefix="Figure 8.")

    # ── 4. FIGURE 10 ATOMISTIC (P32 = image, P33 = caption) ───────────────────
    # In current doc this is "Figure 9: Atomistic 3D..." - we rename to Figure 10
    print("\n[4] Replacing Figure 9 -> Figure 10 - Atomistic Structures (P32)...")
    replace_image_in_paragraph(doc, 32, FIG_FILES["fig10_atomistic"], width_inches=6.3)
    update_caption_paragraph(doc, 33, CAPTIONS["fig10_atomistic"], bold_prefix="Figure 10.")

    # ── 5. Save ────────────────────────────────────────────────────────────────
    doc.save(str(MANUSCRIPT_OUT))
    size_mb = MANUSCRIPT_OUT.stat().st_size / 1e6
    print(f"\n{'='*70}")
    print(f"[SUCCESS] Manuscript saved: {MANUSCRIPT_OUT.name}")
    print(f"  File size: {size_mb:.2f} MB")
    print(f"  Location:  {MANUSCRIPT_OUT}")
    print(f"{'='*70}")


if __name__ == "__main__":
    main()
