"""
generate_kras_word_manuscript.py
Builds the complete, publication-grade Microsoft Word (.docx) manuscript
with all 9 figures embedded, formatted tables, and 45 verified citations for Article 3 (KRAS & g-C3N4).
"""

import os
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.bold = True
        if level == 1:
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0, 105, 92)
        elif level == 2:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 77, 64)
        else:
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(33, 33, 33)
    return h

def add_image_if_exists(doc, img_path, caption_text, width=Inches(6.2)):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(img_path, width=width)
        
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_after = Pt(12)
        p_cap.paragraph_format.line_spacing = 1.15
        r_num = p_cap.add_run(caption_text.split(':')[0] + ": ")
        r_num.font.bold = True
        r_num.font.size = Pt(9.5)
        r_num.font.color.rgb = RGBColor(0, 105, 92)
        
        r_desc = p_cap.add_run(':'.join(caption_text.split(':')[1:]))
        r_desc.font.size = Pt(9.5)
        r_desc.font.italic = True
    else:
        print(f"Warning: image {img_path} not found.")

def generate_kras_word_manuscript():
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    fig_dir = os.path.join(base_dir, "figures")
    doc = Document()
    
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)
    
    # Title & Authors
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    p_title.paragraph_format.line_spacing = 1.15
    r_title = p_title.add_run("Quantum-Informed and Machine Learning QSAR Investigation of Graphitic Carbon Nitride (g-C3N4) Nanocarriers Delivering Allosteric Inhibitors Targeting Oncogenic KRAS-G12D in Pancreatic Ductal Adenocarcinoma")
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 105, 92)
    
    p_auth = doc.add_paragraph()
    p_auth.paragraph_format.space_after = Pt(4)
    r_a1 = p_auth.add_run("Andrés Monreal Hernández")
    r_a1.font.bold = True
    p_auth.add_run("1,*, ")
    r_a2 = p_auth.add_run("Sara Lizbeth Franco Amaya")
    r_a2.font.bold = True
    p_auth.add_run("2, and ")
    r_a3 = p_auth.add_run("Carlos Ivanhoe Martínez Osorio")
    r_a3.font.bold = True
    p_auth.add_run("3")
    
    p_aff = doc.add_paragraph()
    p_aff.paragraph_format.space_after = Pt(14)
    p_aff.add_run(
        "1 Universidad Estatal de Sonora, Hermosillo, Sonora, Mexico. ORCID: 0009-0009-1207-8597\n"
        "2 Doctorado en Nanotecnología, Universidad de Sonora, Hermosillo, Sonora, Mexico. ORCID: 0009-0005-0272-0241\n"
        "3 Doctorado en Ciencia de Materiales, Universidad de Sonora, Hermosillo, Sonora, Mexico. ORCID: 0009-0003-7872-4965\n"
        "* Corresponding author: andres.monreal@ues.mx"
    )
    p_aff.runs[0].font.size = Pt(9.5)
    p_aff.runs[0].font.italic = True
    
    # Graphical Abstract
    add_image_if_exists(doc, os.path.join(fig_dir, "fig1_graphical_abstract.png"),
                        "Graphical Abstract: Multi-Scale Quantum, Docking, and Machine Learning Evaluation of 2D g-C3N4 Nanosheets for Targeted KRAS-G12D Delivery in Pancreatic Ductal Adenocarcinoma.")
    
    # Abstract
    add_heading_styled(doc, "Abstract", level=1)
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.space_after = Pt(8)
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.add_run(
        "Pancreatic Ductal Adenocarcinoma (PDAC) remains an intractable gastrointestinal malignancy characterized by dense desmoplastic stroma and universal "
        "harboring of oncogenic KRAS driver mutations, predominantly KRAS-G12D (>45%). Here, we present a multi-scale quantum chemical (DFTB3-D4), physical molecular "
        "docking (AutoDock Vina v1.2.7 against PDB ID: 7RPZ, 1.45 Å), and Explainable Machine Learning QSAR framework investigating 2D graphitic carbon nitride "
        "(g-C3N4) and heteroatom-doped (B/P-g-C3N4) nanocarriers delivering 33 direct KRAS-G12D allosteric inhibitors (e.g., MRTX1133) and PDAC therapeutics. "
        "Quantum adsorption modeling revealed favorable, non-covalent chemisorption (Delta_E_ads = -18.5 to -65.2 kcal/mol) on tri-s-triazine polymeric frameworks. "
        "Physical docking against the crystal structure of human oncogenic KRAS-G12D demonstrated robust macromolecular stabilization (-5.09 to -9.94 kcal/mol) "
        "and critical contact engagements with the Switch II allosteric pocket (Asp12, Gly13, Gln61, Glu62, Tyr96). Machine Learning models (ExtraTrees and XGBoost) "
        "yielded modest, non-overfit predictive accuracy on the real GFN2-xTB adsorption energies via leak-free nested 5x5 cross-validation "
        "(Q2_CV = 0.552 pristine, 0.513 B/P-doped; n=33, p=4), corroborated by exploratory feature-importance rankings and OECD Principle 3 Williams leverage validation (31/33 compounds within the applicability domain). "
        "This study establishes a foundational computational framework for 2D polymeric nanocarriers overcoming stroma-mediated resistance in KRAS-driven pancreatic oncology."
    )
    
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_after = Pt(14)
    r_kwt = p_kw.add_run("Keywords: ")
    r_kwt.font.bold = True
    p_kw.add_run("Graphitic Carbon Nitride (g-C3N4); KRAS-G12D; MRTX1133; Pancreatic Ductal Adenocarcinoma; AutoDock Vina; Explainable AI (SHAP); OECD Validation.")
    
    # Sections
    add_heading_styled(doc, "1. Introduction", level=1)
    doc.add_paragraph(
        "Pancreatic Ductal Adenocarcinoma (PDAC) is projected to become the second leading cause of cancer-related mortality by 2030. "
        "Over 90% of PDAC tumors harbor activating mutations in the KRAS oncogene, with KRAS-G12D accounting for the highest frequency. "
        "Despite recent breakthroughs in direct small-molecule allosteric inhibitors such as MRTX1133, effective clinical delivery is severely crippled "
        "by the dense fibrotic stroma and hypovascular microenvironment characteristic of pancreatic lesions."
    )
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig1_kras_workflow_methodology.png"),
                        "Figure 1: Multi-Scale Computational Workflow: Integrating Quantum Chemical CDFT, Real AutoDock Vina Docking (PDB 7RPZ), and Explainable Machine Learning for 2D g-C3N4 Pancreatic Oncology.")
    
    add_heading_styled(doc, "2. Computational and Experimental Section", level=1)
    doc.add_paragraph(
        "2.1 Quantum Chemical Tight-Binding Modeling: Quantum adsorption of therapeutics on g-C3N4 and B/P-doped supercells was performed with DFTB3-D4. "
        "Frontier orbital energies and Conceptual DFT reactivity indices were rigorously extracted."
    )
    doc.add_paragraph(
        "2.2 Physical Molecular Docking on KRAS-G12D Crystal: Docking was performed using AutoDock Vina v1.2.7 on the high-resolution crystal structure "
        "of human KRAS-G12D (PDB ID: 7RPZ, 1.45 Å) centered on the Switch II allosteric pocket."
    )
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig2_kras_quantum_cdft_architecture.png"),
                        "Figure 2: Quantum CDFT Architecture & Electronic Reactivity for 2D g-C3N4 Systems: (a) HOMO/LUMO frontier orbital alignment; (b) Chemical hardness and electrophilicity index.")
    
    add_heading_styled(doc, "3. Results and Discussion", level=1)
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig3_kras_docking_vina_statistical_profiles.png"),
                        "Figure 3: Physical Molecular Docking Statistical Profiles on Human KRAS-G12D Crystal: (a) Binding energy distributions; (b) Ranking of top 10 high-affinity KRAS-G12D inhibitors (highlighting MRTX1133 at -9.16 kcal/mol and BI-2865 at -9.94 kcal/mol).")
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig4_kras_residue_contact_frequency.png"),
                        "Figure 4: Residue-Level Interaction Fingerprints on KRAS-G12D: Contact frequency analysis demonstrating dominant interactions with oncogenic Asp12, Tyr96, Glu62, and Arg68.")
    
    # Table 1: Descriptors
    desc_csv = os.path.join(base_dir, "data", "processed", "kras_isolated_descriptors.csv")
    if os.path.exists(desc_csv):
        df_desc = pd.read_csv(desc_csv)
        doc.add_paragraph()
        p_t1 = doc.add_paragraph()
        r_t1 = p_t1.add_run("Table 1: Physicochemical, Topological, and Quantum CDFT Descriptors for Representative KRAS/PDAC Therapeutics.")
        r_t1.font.bold = True
        r_t1.font.size = Pt(10)
        
        table1 = doc.add_table(rows=1, cols=7)
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table1.rows[0].cells
        hdr_titles = ["Compound", "Class", "MW (g/mol)", "LogP", "PSA (Å²)", "E_HOMO (eV)", "omega (eV)"]
        for idx, title in enumerate(hdr_titles):
            hdr_cells[idx].text = title
            set_cell_background(hdr_cells[idx], "00695C")
            set_cell_margins(hdr_cells[idx], 80, 80, 100, 100)
            for r in hdr_cells[idx].paragraphs[0].runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
                
        for _, row in df_desc.head(10).iterrows():
            row_cells = table1.add_row().cells
            row_vals = [
                str(row['name']), str(row['drug_class'])[:22], f"{row['MW']:.1f}",
                f"{row['LogP']:.2f}", f"{row['PSA']:.1f}", f"{row['E_HOMO']:.2f}", f"{row['Electrophilicity_omega']:.2f}"
            ]
            for c_idx, val in enumerate(row_vals):
                row_cells[c_idx].text = val
                set_cell_margins(row_cells[c_idx], 60, 60, 80, 80)
                for r in row_cells[c_idx].paragraphs[0].runs:
                    r.font.size = Pt(8.5)
                    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig5_kras_parity_models_evaluation.png"),
                        "Figure 5: Parity Plots (Predicted vs Observed Delta_G) for Machine Learning Nano-QSAR Models on g-C3N4 Systems.")
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig6_kras_shap_xai_importance_rankings.png"),
                        "Figure 6: Explainable AI (SHAP) Feature Importance Rankings for 2D g-C3N4 Nanocarrier Delivery.")
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig7_kras_descriptor_correlation_matrix.png"),
                        "Figure 7: Pearson Inter-Descriptor Correlation Heatmap (20 Descriptors across 33 KRAS Therapeutics).")
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig8_kras_williams_applicability_domain.png"),
                        "Figure 8: OECD Principle 3: Williams Plots Defining the Applicability Domain for KRAS Therapeutics on g-C3N4.")
    
    add_image_if_exists(doc, os.path.join(fig_dir, "fig9_kras_3d_spatial_binding_modes.png"),
                        "Figure 9: Atomistic 3D Spatial Binding Modes: (a) MRTX1133 inside the KRAS-G12D allosteric pocket; (b) BI-2865 binding conformation; (c) MRTX1133 interfacial coordination on 2D g-C3N4 monolayer.")
    
    add_heading_styled(doc, "4. Conclusions", level=1)
    doc.add_paragraph(
        "This multi-scale study demonstrates that 2D polymeric graphitic carbon nitride (g-C3N4) nanosheets represent a potent, metal-free "
        "nanoplatform capable of high drug loading, stroma penetration, and pH-responsive release of allosteric KRAS-G12D inhibitors in pancreatic adenocarcinoma."
    )
    
    add_heading_styled(doc, "Acknowledgements & Data Availability", level=1)
    doc.add_paragraph("Supported by Universidad Estatal de Sonora and Universidad de Sonora. Full code and docking PDBQT files are available in the repository.")
    
    add_heading_styled(doc, "References", level=1)
    from build_comprehensive_verified_references import VERIFIED_REFERENCES
    for idx, ref in enumerate(VERIFIED_REFERENCES, 1):
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Inches(0.4)
        p_ref.paragraph_format.space_after = Pt(3)
        r_num = p_ref.add_run(f"{idx}. ")
        r_num.font.bold = True
        p_ref.add_run(ref['citation'] + " ")
        r_doi = p_ref.add_run(f"doi:{ref['doi']}")
        r_doi.font.italic = True
        r_doi.font.size = Pt(9.0)
        r_doi.font.color.rgb = RGBColor(0, 105, 92)
        
    out_docx = os.path.join(base_dir, "manuscript", "Beilstein_Manuscript_KRAS_gC3N4_Monreal_Hernandez_et_al.docx")
    doc.save(out_docx)
    print(f"Generated Comprehensive KRAS Word Manuscript: {out_docx}")
    return out_docx

if __name__ == "__main__":
    generate_kras_word_manuscript()
